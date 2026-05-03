"""
D.O. (Demi-Official) Letter Generator
---------------------------------------
Generates Indian government-style D.O. letters for grievance escalation.

Flow
----
1.  select_template(action_type, category)  →  template filename
2.  Claude Sonnet produces letter content fields as strict JSON
3.  Jinja2 renders the chosen HTML template with those fields
4.  Returns { html, do_number, letter_type, letter_fields }

D.O. Number format:  {do_prefix}/{year}-{counter:04d}
  e.g.  OFF/25/2026-0003

The counter is the count of existing letters_log rows for this office + 1
(simple, restartable, no race condition for low-volume usage).
"""

import io
import json
import logging
import os
from datetime import datetime, timezone
from pathlib import Path

from anthropic import AsyncAnthropic
from jinja2 import Environment, FileSystemLoader, select_autoescape

from app.core.config import settings
from app.core.database import get_db

logger = logging.getLogger(__name__)

# ── Template directory ────────────────────────────────────────────────────────
_TEMPLATES_DIR = Path(__file__).parent.parent / "templates" / "letters"

_jinja_env = Environment(
    loader=FileSystemLoader(str(_TEMPLATES_DIR)),
    autoescape=select_autoescape(["html"]),
)

# ── Template selection lookup ─────────────────────────────────────────────────
# Keys: (action_type, category)  →  template filename
# Add more entries here to support additional templates without touching logic.
_TEMPLATE_MAP: dict[tuple[str, str], str] = {
    # Infrastructure / field-visit complaints → inspection request variant
    ("field_visit",       "infrastructure"): "do_inspection_request.html",
    ("escalate_to_dept",  "infrastructure"): "do_inspection_request.html",
    ("call_official",     "infrastructure"): "do_inspection_request.html",
    # Railway / transport complaints → railway quota letter
    ("railway_quota",     "*"):              "railway_quota.html",
    ("escalate_to_dept",  "transport"):      "railway_quota.html",
    ("forward_to_dept",   "transport"):      "railway_quota.html",
}
_DEFAULT_TEMPLATE = "do_standard.html"

# Keywords that indicate a railway quota/ticket request (checked against raw_text + summary)
_RAILWAY_KEYWORDS = {
    "pnr", "train", "railway", "reservation", "ticket", "berth", "irctc",
    "quota", "coach", "seat", "tatkal", "waitlist", "waiting list", "rac",
    "रेलवे", "टिकट", "ट्रेन", "रिजर्वेशन",
}


def _is_railway_request(grievance: dict) -> bool:
    """Return True if the grievance text strongly suggests a railway quota request."""
    text = " ".join([
        (grievance.get("summary") or ""),
        (grievance.get("raw_text") or ""),
        (grievance.get("description") or ""),
    ]).lower()
    return any(kw in text for kw in _RAILWAY_KEYWORDS)


def select_template(action_type: str, category: str, grievance: dict | None = None) -> str:
    """
    Return the Jinja2 template filename for the given action_type + category.

    Lookup order:
      1. Railway keyword auto-detection (if grievance dict provided)
      2. Exact (action_type, category) match
      3. (action_type, '*') wildcard
      4. Default: do_standard.html
    """
    if grievance and _is_railway_request(grievance):
        return "railway_quota.html"
    key = (action_type.lower().strip(), category.lower().strip())
    if key in _TEMPLATE_MAP:
        return _TEMPLATE_MAP[key]
    wildcard_key = (action_type.lower().strip(), "*")
    return _TEMPLATE_MAP.get(wildcard_key, _DEFAULT_TEMPLATE)


# ── Claude prompt ─────────────────────────────────────────────────────────────
_SYSTEM_PROMPT = """\
You are a senior Private Secretary drafting a D.O. (Demi-Official) letter on behalf of an Indian elected \
representative.

D.O. letters in Indian government are semi-formal — less rigid than an official demi-official letter but \
more authoritative than a WhatsApp message. They are addressed personally to a named officer and seek \
specific action.

TONE RULES:
- Formal, measured, respectful — not WhatsApp-style
- Third-person reference to the representative: "Hon'ble [title]..." or "the Hon'ble Member..."
- Opening: "As desired by Hon'ble [rep_title], I am directed to bring to your kind notice..."
- Never use bullet points in the body — flowing paragraphs only
- Cite the grievance reference ID and location in the body
- End with a polite but firm request for action within a specified timeframe

REQUIRED JSON SCHEMA — respond ONLY with valid JSON, no prose, no markdown fences:
{
  "subject": "<Subject line — concise, formal, referencing the grievance topic and location>",
  "salutation": "Sir,",
  "opening_para": "<First paragraph: cite the grievance ID, describe the issue, mention the citizen/location>",
  "body_paras": ["<Second paragraph: context, severity, why action is needed now>",
                 "<Third paragraph: specific ask — what department must do, by when>"],
  "closing_note": "<One sentence: request for an early action report / acknowledgement>",
  "addressee_name": "<Full name of the officer to address — infer from target_dept or use 'The Concerned Officer'>",
  "addressee_designation": "<Official title/rank of addressee>",
  "addressee_org": "<Department / organisation name>",
  "addressee_address_lines": ["<Line 1>", "<City — PIN>"],
  "addressee_email": "<official email if inferable, else leave empty string>"
}

Do NOT add any fields beyond this schema. Values must be plain strings (no markdown inside values).
"""

_RAILWAY_SYSTEM_PROMPT = """\
You are a senior Private Secretary extracting travel details from a railway reservation request submitted by a \
constituent, and drafting a one-sentence polite opening line for the MP/MLA's railway quota letter.

Your job is to:
1. Extract every travel detail that appears in the grievance text — PNR number, train number, date of journey, \
   travel class, from/to station, passenger name, and passenger contact number.
2. If a detail is not mentioned, return an empty string for that field.
3. Draft a single formal opening sentence for the MP/MLA to send to the Chief Reservation Supervisor.

REQUIRED JSON SCHEMA — respond ONLY with valid JSON, no prose, no markdown fences:
{
  "opening_line": "<Single formal sentence: 'I shall be highly grateful to you, if you could kindly make an arrangement for ticket confirmation / reservation for my constituent / guest, whose details are given below:'>",
  "pnr_number":        "<PNR number, or empty string>",
  "train_number":      "<Train number, or empty string>",
  "travel_date":       "<Date of journey in DD.MM.YYYY format, or empty string>",
  "travel_class":      "<Class code e.g. 2A / 3A / SL / CC, or empty string>",
  "from_station":      "<Departure station name, or empty string>",
  "to_station":        "<Destination station name, or empty string>",
  "passenger_name":    "<Full name of passenger, or empty string>",
  "passenger_contact": "<Mobile number of passenger, or empty string>",
  "addressee_name":    "OSD to Hon'ble Minister of Railways",
  "addressee_designation": "",
  "addressee_org":     "Government of India",
  "addressee_address_lines": ["Rail Bhawan, Raisina Road", "New Delhi - 110001"]
}

The addressee is ALWAYS "OSD to Hon'ble Minister of Railways, Government of India, Rail Bhawan, Raisina Road, New Delhi - 110001". Do NOT change these values regardless of what the grievance says.
Do NOT add any fields beyond this schema.
"""


async def _call_claude_for_railway_fields(grievance: dict, office: dict) -> dict:
    """Call Claude to extract railway travel details from the grievance text."""
    client = AsyncAnthropic(api_key=settings.ANTHROPIC_API_KEY)

    profile       = office.get("letter_profile") or {}
    grievance_ref = grievance.get("grievance_id", "N/A")
    raw_text      = grievance.get("raw_text") or grievance.get("description") or ""
    summary       = grievance.get("summary", "")
    citizen_name  = grievance.get("citizen_name", "")

    user_content = f"""
GRIEVANCE REFERENCE: {grievance_ref}
CITIZEN NAME: {citizen_name}

GRIEVANCE TEXT:
{raw_text or summary}

Extract the railway travel details and draft the opening line now.
"""

    response = await client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=600,
        system=_RAILWAY_SYSTEM_PROMPT,
        messages=[{"role": "user", "content": user_content}],
    )

    raw = response.content[0].text.strip()
    start, end = raw.find("{"), raw.rfind("}")
    if start != -1 and end != -1:
        raw = raw[start : end + 1]

    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        logger.error("Railway letter generator returned non-JSON: %s", raw)
        return {
            "opening_line": (
                "I shall be highly grateful to you, if you could kindly make an arrangement "
                "for ticket confirmation / reservation for my constituent, whose details are given below:"
            ),
            "pnr_number": "",
            "train_number": "",
            "travel_date": "",
            "travel_class": "",
            "from_station": "",
            "to_station": "",
            "passenger_name": citizen_name,
            "passenger_contact": grievance.get("citizen_contact", ""),
            "addressee_name": "OSD to Hon'ble Minister of Railways",
            "addressee_designation": "",
            "addressee_org": "Government of India",
            "addressee_address_lines": ["Rail Bhawan, Raisina Road", "New Delhi - 110001"],
        }


async def _call_claude_for_letter_fields(grievance: dict, office: dict, rep_title: str) -> dict:
    """Call Claude Sonnet to produce the letter content fields."""
    client = AsyncAnthropic(api_key=settings.ANTHROPIC_API_KEY)

    profile = office.get("letter_profile") or {}
    target_dept = (
        grievance.get("suggested_action", "")
        or profile.get("rep_designation", "the representative")
    )

    # Build target_dept from action advisor result if present in grievance
    # (stored as suggested_action text; we try to extract the dept name)
    action_type   = grievance.get("_action_type", "escalate_to_dept")
    category      = grievance.get("category", "others")
    location      = grievance.get("location_text") or "Not specified"
    summary       = grievance.get("summary", "")
    grievance_ref = grievance.get("grievance_id", "N/A")
    urgency       = grievance.get("urgency", "medium").upper()

    user_content = f"""
GRIEVANCE DETAILS:
  Reference ID   : {grievance_ref}
  Category       : {category}
  Urgency        : {urgency}
  Location       : {location}
  Summary        : {summary}
  Action type    : {action_type}
  Target dept    : {grievance.get("_target_dept", "District Collector")}

REPRESENTATIVE PROFILE:
  Title          : {rep_title}
  Designation    : {profile.get("rep_designation", "[Designation]")}

Draft the D.O. letter content JSON now.
"""

    response = await client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=900,
        system=_SYSTEM_PROMPT,
        messages=[{"role": "user", "content": user_content}],
    )

    raw = response.content[0].text.strip()
    # Strip markdown fences if Claude added them despite instructions
    start, end = raw.find("{"), raw.rfind("}")
    if start != -1 and end != -1:
        raw = raw[start : end + 1]

    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        logger.error("Letter generator Claude returned non-JSON: %s", raw)
        # Graceful fallback — minimal valid letter fields
        return {
            "subject": f"Re: Grievance {grievance_ref} — {category.replace('_', ' ').title()}",
            "salutation": "Sir,",
            "opening_para": (
                f"As desired by Hon'ble {rep_title}, I am directed to bring to your kind notice "
                f"grievance reference {grievance_ref} received from a constituent regarding "
                f"{summary or category.replace('_', ' ')} at {location}."
            ),
            "body_paras": [
                "The matter requires urgent attention and early resolution to address the "
                "concern raised by the constituent."
            ],
            "closing_note": "An early action report may kindly be sent to this office.",
            "addressee_name": "The Concerned Officer",
            "addressee_designation": grievance.get("_target_dept", "Department Head"),
            "addressee_org": grievance.get("_target_dept", "Concerned Department"),
            "addressee_address_lines": ["Office Address", "City"],
            "addressee_email": "",
        }


def _build_do_number(office: dict, counter: int) -> str:
    """
    Compose the D.O. number.
    Format: {do_prefix}/{year}-{counter:04d}
    e.g.   OFF/25/2026-0003
    """
    profile   = office.get("letter_profile") or {}
    do_prefix = profile.get("do_prefix", "OFF")
    year      = datetime.now(tz=timezone.utc).year
    return f"{do_prefix}/{year}-{counter:04d}"


def _get_letter_counter(office_id: str) -> int:
    """Return count of existing letters_log rows for this office + 1."""
    try:
        db  = get_db()
        res = (
            db.table("letters_log")
            .select("id", count="exact")
            .eq("office_id", office_id)
            .execute()
        )
        existing = res.count if hasattr(res, "count") and res.count is not None else len(res.data or [])
        return existing + 1
    except Exception:
        logger.exception("Failed to count letters_log for office %s — using 1", office_id)
        return 1


def _log_letter(
    office_id: str,
    grievance_id_uuid: str | None,
    do_number: str,
    letter_type: str,
    addressee_name: str,
    html_content: str,
) -> None:
    """Persist the generated letter to letters_log for audit trail."""
    try:
        db = get_db()
        db.table("letters_log").insert({
            "office_id":      office_id,
            "grievance_id":   grievance_id_uuid,
            "do_number":      do_number,
            "letter_type":    letter_type,
            "addressee_name": addressee_name,
            "html_content":   html_content,
            "generated_at":   datetime.now(tz=timezone.utc).isoformat(),
        }).execute()
    except Exception:
        logger.exception("Failed to log letter (do_number=%s) — continuing anyway", do_number)


def _generate_pdf(html: str) -> bytes:
    """Convert rendered HTML to PDF bytes using WeasyPrint."""
    try:
        from weasyprint import HTML
        return HTML(string=html).write_pdf()
    except Exception:
        logger.exception("WeasyPrint PDF generation failed")
        raise


def _generate_docx(context: dict, letter_type: str) -> bytes:
    """
    Build a clean, editable Word document from the letter context.
    Uses python-docx directly — does not parse HTML.
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt, RGBColor

    doc = Document()

    # ── Page setup: A4 ───────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2.5)

    # Accent colour (maroon for D.O. letters, green for railway)
    is_railway = (letter_type == "railway_quota")
    accent = RGBColor(0x1a, 0x6b, 0x1a) if is_railway else RGBColor(0x8B, 0x00, 0x00)

    def _para(text: str, bold=False, italic=False, size=11,
              align=WD_ALIGN_PARAGRAPH.LEFT, color=None, space_before=0, space_after=6):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.size  = Pt(size)
        run.font.color.rgb = color or RGBColor(0, 0, 0)
        return p

    def _rule():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(8)
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:color"), f"{accent[0]:02X}{accent[1]:02X}{accent[2]:02X}")
        pBdr.append(bottom)
        pPr.append(pBdr)

    # ── Header ────────────────────────────────────────────────────────────────
    _para(context.get("rep_name_hindi", ""), bold=True, size=14, color=accent)
    _para(context.get("rep_designation", ""), bold=True, size=10, space_after=2)

    memberships = context.get("committee_memberships") or []
    for m in memberships:
        _para(f"• {m}", size=9, space_after=1)

    _para(context.get("office_address", ""), size=9, space_after=2)
    phone = context.get("office_phone", "")
    mobile = context.get("office_mobile", "")
    if phone:
        _para(f"Phone: {phone}", size=9, space_after=1)
    if mobile:
        _para(f"Mobile: {mobile}", size=9, space_after=1)

    _rule()

    # ── D.O. Number + Date ────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(12)
    r1 = p.add_run(f"D.O. No. {context.get('do_number', '')}    ")
    r1.font.size = Pt(10)
    r2 = p.add_run(f"Date: {context.get('letter_date', '')}")
    r2.font.size = Pt(10)
    r2.italic = True

    # ── Addressee (railway: top; D.O.: below subject) ─────────────────────────
    def _write_addressee():
        _para(context.get("addressee_name", ""), bold=True, size=11, space_after=2)
        if context.get("addressee_designation"):
            _para(context["addressee_designation"], size=11, space_after=2)
        if context.get("addressee_org"):
            _para(context["addressee_org"], size=11, space_after=2)
        for line in (context.get("addressee_address_lines") or []):
            _para(line, size=11, space_after=2)
        doc.add_paragraph()

    if is_railway:
        _write_addressee()

    # ── Subject (D.O. letters only) ───────────────────────────────────────────
    subject = context.get("subject", "")
    if subject and not is_railway:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        r = p.add_run(f"Sub: {subject}")
        r.bold = True
        r.font.size = Pt(11)
        r.font.underline = True

        _write_addressee()

    # ── Salutation / Opening ──────────────────────────────────────────────────
    if is_railway:
        _para(context.get("opening_line", ""), size=11, space_after=10)
    else:
        _para(context.get("salutation", "Sir,"), size=11, space_after=8)
        opening = context.get("opening_para", "")
        if opening:
            p = doc.add_paragraph(opening)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Cm(1)
            p.paragraph_format.space_after = Pt(8)

        for para_text in (context.get("body_paras") or []):
            p = doc.add_paragraph(para_text)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Cm(1)
            p.paragraph_format.space_after = Pt(8)

        closing = context.get("closing_note", "")
        if closing:
            p = doc.add_paragraph(closing)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Cm(1)
            p.paragraph_format.space_after = Pt(16)

    # ── Railway details table ─────────────────────────────────────────────────
    if is_railway:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        table = doc.add_table(rows=7, cols=3)
        try:
            table.style = "Table Grid"
        except KeyError:
            pass  # style not in default template — borders added manually below
        table.autofit = True

        rows_data = [
            ("1.", "PNR No.",     context.get("pnr_number", "")),
            ("2.", "Train No.",   context.get("train_number", "")),
            ("3.", "Date",        context.get("travel_date", "")),
            ("4.", "Class",       context.get("travel_class", "")),
            ("5.", "Destination", f"From {context.get('from_station','')} To {context.get('to_station','')}"),
            ("6.", "Name",        context.get("passenger_name", "")),
            ("7.", "Contact No.", context.get("passenger_contact", "")),
        ]
        for i, (num, label, value) in enumerate(rows_data):
            row = table.rows[i]
            row.cells[0].text = num
            row.cells[1].text = label
            row.cells[2].text = value
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(11)

        doc.add_paragraph()

    # ── Sign-off ──────────────────────────────────────────────────────────────
    _para("With Regards," if is_railway else "Yours faithfully,", size=11, space_before=16, space_after=36)
    _para(f"({context.get('sender_name', '')})", size=11, space_after=2)
    if is_railway and context.get("ic_number"):
        _para(f"IC No. {context['ic_number']}", size=10, space_after=2)

    # ── Footer ────────────────────────────────────────────────────────────────
    footer_text = context.get("office_footer", "")
    if footer_text:
        _rule()
        _para(footer_text, size=8, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x33, 0x33, 0x33))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


async def generate_do_letter(
    grievance: dict,
    office: dict,
    letter_type_override: str | None = None,
) -> dict:
    """
    Main entry point: produce a complete D.O. letter for a grievance.

    Parameters
    ----------
    grievance : dict
        Grievance row from Supabase.
    office : dict
        Office row from Supabase (must include letter_profile JSONB).
    letter_type_override : str | None
        Force a specific template stem (e.g. "railway_quota").
        If None, auto-detected from action_type + category + keywords.

    Returns
    -------
    dict with keys:
        html         : rendered HTML string (ready for browser / download)
        do_number    : e.g. "OFF/25/2026-0003"
        letter_type  : template filename stem (e.g. "do_standard")
        letter_fields: raw dict returned by Claude
    """
    profile  = office.get("letter_profile") or {}
    action_type = grievance.get("_action_type", "escalate_to_dept")
    category    = grievance.get("category", "others")

    # ── 1. Pick template ──────────────────────────────────────────────────────
    if letter_type_override:
        template_file = f"{letter_type_override}.html"
    else:
        template_file = select_template(action_type, category, grievance)
    letter_type = template_file.replace(".html", "")

    # ── 2. Build rep title (safe fallbacks for missing profile fields) ────────
    rep_title = (
        profile.get("rep_full_title")
        or profile.get("rep_designation")
        or "[REPRESENTATIVE TITLE]"
    )

    # ── 3. Call Claude for letter content (branch on template type) ───────────
    is_railway = (letter_type == "railway_quota")
    if is_railway:
        fields = await _call_claude_for_railway_fields(grievance, office)
    else:
        fields = await _call_claude_for_letter_fields(grievance, office, rep_title)

    # ── 4. Compute D.O. number ────────────────────────────────────────────────
    office_id = str(office.get("id", ""))
    counter   = _get_letter_counter(office_id)
    do_number = _build_do_number(office, counter)

    # ── 5. Assemble template context ──────────────────────────────────────────
    today_str = datetime.now(tz=timezone.utc).strftime("%d %B %Y")

    # Parse office_address into a list of lines for railway template
    raw_address = profile.get("office_address", "[Office Address]")
    office_address_lines = (
        [l.strip() for l in raw_address.split(",") if l.strip()]
        if isinstance(raw_address, str)
        else raw_address
    )

    context = {
        # Letterhead — rep / sender side
        "rep_name":          profile.get("rep_name", "[REPRESENTATIVE NAME]"),
        "rep_name_hindi":    profile.get("rep_name_hindi", "[प्रतिनिधि का नाम]"),
        "rep_designation":   profile.get("rep_designation", "[Designation]"),
        "rep_full_title":    rep_title,
        "sender_name":       profile.get("sender_name", "[Sender Name]"),
        "sender_role_english": profile.get(
            "sender_role_english", "PRIVATE SECRETARY TO THE REPRESENTATIVE"
        ),
        "sender_role_hindi": profile.get("sender_role_hindi", "[हिंदी में पदनाम]"),
        "office_address":    raw_address,
        "office_address_lines": office_address_lines,
        "office_phone":      profile.get("office_phone", ""),
        "office_mobile":     profile.get("office_mobile", ""),
        "office_fax":        profile.get("office_fax", ""),
        "office_email":      profile.get("office_email", ""),
        "office_footer":     profile.get("office_footer", ""),
        # Committee memberships (for railway quota header)
        "committee_memberships": profile.get("committee_memberships", []),
        # IC number (for railway quota sign-off)
        "ic_number":         profile.get("ic_number", ""),
        # Letter metadata
        "do_number":         do_number,
        "letter_date":       today_str,
        # Letter content from Claude
        "subject":           fields.get("subject", ""),
        "salutation":        fields.get("salutation", "Sir,"),
        "opening_para":      fields.get("opening_para", ""),
        "opening_line":      fields.get("opening_line", ""),
        "body_paras":        fields.get("body_paras", []),
        "closing_note":      fields.get("closing_note", ""),
        # Inspection-specific (used only in do_inspection_request.html)
        "inspection_location": grievance.get("location_text") or "[Location]",
        # Railway-specific fields
        "pnr_number":        fields.get("pnr_number", ""),
        "train_number":      fields.get("train_number", ""),
        "travel_date":       fields.get("travel_date", ""),
        "travel_class":      fields.get("travel_class", ""),
        "from_station":      fields.get("from_station", ""),
        "to_station":        fields.get("to_station", ""),
        "passenger_name":    fields.get("passenger_name", ""),
        "passenger_contact": fields.get("passenger_contact", ""),
        # Addressee block
        "addressee_name":    fields.get("addressee_name", "The Concerned Officer"),
        "addressee_designation": fields.get("addressee_designation", ""),
        "addressee_org":     fields.get("addressee_org", ""),
        "addressee_address_lines": fields.get("addressee_address_lines", []),
        "addressee_email":   fields.get("addressee_email", ""),
        # Misc
        "grievance_id":      grievance.get("grievance_id", ""),
        "urgency":           grievance.get("urgency", "medium"),
    }

    # ── 6. Render Jinja2 template ─────────────────────────────────────────────
    try:
        tmpl = _jinja_env.get_template(template_file)
        html = tmpl.render(**context)
    except Exception:
        logger.exception("Jinja2 render failed for template %s", template_file)
        raise

    # ── 7. Generate PDF and DOCX ──────────────────────────────────────────────
    import base64

    try:
        pdf_bytes  = _generate_pdf(html)
        pdf_b64    = base64.b64encode(pdf_bytes).decode()
    except Exception:
        logger.exception("PDF generation failed — continuing without PDF")
        pdf_b64 = ""

    try:
        docx_bytes = _generate_docx(context, letter_type)
        docx_b64   = base64.b64encode(docx_bytes).decode()
    except Exception:
        logger.exception("DOCX generation failed — continuing without DOCX")
        docx_b64 = ""

    # ── 8. Log to letters_log ─────────────────────────────────────────────────
    _log_letter(
        office_id=office_id,
        grievance_id_uuid=grievance.get("id"),
        do_number=do_number,
        letter_type=letter_type,
        addressee_name=fields.get("addressee_name", ""),
        html_content=html,
    )

    return {
        "html":          html,
        "pdf_b64":       pdf_b64,
        "docx_b64":      docx_b64,
        "do_number":     do_number,
        "letter_type":   letter_type,
        "letter_fields": fields,
    }


async def generate_birthday_letter(citizen: dict, office: dict) -> dict:
    """
    Generate a birthday wishes letter for a citizen.

    Parameters
    ----------
    citizen : dict
        Must contain: name, dob (date object or ISO string), salutation,
        designation (optional), address_lines (list, optional).
    office : dict
        Office row from Supabase (must include letter_profile JSONB).

    Returns
    -------
    dict with keys: html, pdf_b64, docx_b64, letter_type
    """
    import base64

    profile  = office.get("letter_profile") or {}
    today    = datetime.now(tz=timezone.utc)
    today_str = today.strftime("%B %d, %Y")

    # Parse citizen DOB for display
    dob_raw = citizen.get("dob", "")
    try:
        if isinstance(dob_raw, str):
            from datetime import date
            dob_dt = date.fromisoformat(dob_raw[:10])
        else:
            dob_dt = dob_raw
        birthday_date = dob_dt.strftime("%d.%m.%Y")
    except Exception:
        birthday_date = str(dob_raw)

    # Citizen salutation — default "Shri" for unknown gender
    salutation = citizen.get("salutation") or citizen.get("gender_salutation") or "Shri"

    raw_address = profile.get("office_address", "[Office Address]")

    context = {
        # Letterhead
        "rep_name":           profile.get("rep_name", "[REPRESENTATIVE NAME]"),
        "rep_name_hindi":     profile.get("rep_name_hindi", "[प्रतिनिधि का नाम]"),
        "rep_designation":    profile.get("rep_designation", "[Designation]"),
        "sender_role_english": profile.get("sender_role_english", ""),
        "sender_role_hindi":  profile.get("sender_role_hindi", ""),
        "office_address":     raw_address,
        # Letter metadata
        "letter_date":        today_str,
        # Citizen details
        "citizen_name":        citizen.get("name", ""),
        "citizen_salutation":  salutation,
        "citizen_designation": citizen.get("designation", ""),
        "citizen_address_lines": citizen.get("address_lines", []),
        "birthday_date":       birthday_date,
    }

    # Render HTML
    try:
        tmpl = _jinja_env.get_template("birthday_wishes.html")
        html = tmpl.render(**context)
    except Exception:
        logger.exception("Jinja2 render failed for birthday_wishes.html")
        raise

    # Generate PDF and DOCX
    try:
        pdf_bytes = _generate_pdf(html)
        pdf_b64   = base64.b64encode(pdf_bytes).decode()
    except Exception:
        logger.exception("PDF generation failed for birthday letter")
        pdf_b64 = ""

    try:
        docx_bytes = _generate_birthday_docx(context)
        docx_b64   = base64.b64encode(docx_bytes).decode()
    except Exception:
        logger.exception("DOCX generation failed for birthday letter")
        docx_b64 = ""

    # Log to letters_log (no grievance link)
    office_id = str(office.get("id", ""))
    counter   = _get_letter_counter(office_id)
    do_number = _build_do_number(office, counter)

    _log_letter(
        office_id=office_id,
        grievance_id_uuid=None,
        do_number=do_number,
        letter_type="birthday_wishes",
        addressee_name=citizen.get("name", ""),
        html_content=html,
    )

    return {
        "html":        html,
        "pdf_b64":     pdf_b64,
        "docx_b64":    docx_b64,
        "letter_type": "birthday_wishes",
        "do_number":   do_number,
    }


def _generate_birthday_docx(context: dict) -> bytes:
    """Build a clean Word document for a birthday wishes letter."""
    import io
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt, RGBColor

    doc = Document()
    section = doc.sections[0]
    section.page_width    = Cm(21)
    section.page_height   = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2.5)

    maroon = RGBColor(0x80, 0x00, 0x00)
    black  = RGBColor(0x00, 0x00, 0x00)

    def _p(text, bold=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT,
           color=None, space_before=0, space_after=6, italic=False):
        para = doc.add_paragraph()
        para.alignment = align
        para.paragraph_format.space_before = Pt(space_before)
        para.paragraph_format.space_after  = Pt(space_after)
        run = para.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.size = Pt(size)
        run.font.color.rgb = color or black
        return para

    # Header
    _p(context.get("rep_name_hindi", ""), bold=True, size=13)
    _p(context.get("rep_name", ""),       bold=True, size=11, space_after=2)
    _p(context.get("rep_designation", ""),size=9,  color=RGBColor(0x33,0x33,0x33), space_after=2)
    _p(context.get("office_address", ""), size=9,  color=RGBColor(0x33,0x33,0x33), space_after=8)

    # Separator rule
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    rule_p = doc.add_paragraph()
    rule_p.paragraph_format.space_after = Pt(10)
    pPr = rule_p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "double")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "800000")
    pBdr.append(bottom)
    pPr.append(pBdr)

    # Date (right-aligned)
    _p(context.get("letter_date", ""), size=11, align=WD_ALIGN_PARAGRAPH.RIGHT,
       italic=True, space_before=10, space_after=18)

    # Salutation
    sal = f"Dear {context.get('citizen_salutation','')} {context.get('citizen_name','')},"
    _p(sal, bold=True, size=11, space_after=14)

    # Body paragraphs (indented)
    for text in [
        f"I convey my best wishes on your birthday on {context.get('birthday_date', '')}.",
        "May God bless you with good health and happiness and many more years of "
        "continued and dedicated service to the nation.",
    ]:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.first_line_indent = Cm(1.2)
        para.paragraph_format.space_after = Pt(12)
        run = para.add_run(text)
        run.font.size = Pt(11)

    # Sign-off
    _p("With the best wishes,", size=11, space_before=10, space_after=8)
    _p("Yours sincerely,", size=11, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=48)
    _p(f"({context.get('rep_name', '')})", size=11, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=24)

    # Addressee bottom block
    from docx.oxml import OxmlElement
    rule_p2 = doc.add_paragraph()
    rule_p2.paragraph_format.space_after = Pt(8)
    pPr2 = rule_p2._p.get_or_add_pPr()
    pBdr2 = OxmlElement("w:pBdr")
    top2 = OxmlElement("w:top")
    top2.set(qn("w:val"), "single")
    top2.set(qn("w:sz"), "4")
    top2.set(qn("w:color"), "CCCCCC")
    pBdr2.append(top2)
    pPr2.append(pBdr2)

    citizen_name = f"{context.get('citizen_salutation','')} {context.get('citizen_name','')}".strip()
    _p(citizen_name, bold=True, size=10, space_after=2)
    if context.get("citizen_designation"):
        _p(context["citizen_designation"], size=10, space_after=2)
    for line in (context.get("citizen_address_lines") or []):
        _p(line, size=10, space_after=2)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
